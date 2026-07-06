from dataclasses import dataclass
from datetime import UTC, datetime
from email import message_from_binary_file
from email.message import Message
from email.utils import parsedate_to_datetime
from pathlib import Path
import shlex
import subprocess
import tempfile
import re
import zipfile
from xml.etree import ElementTree

from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.config import settings


@dataclass(frozen=True)
class ExtractedAttachment:
    filename: str
    file_type: str
    payload: bytes
    text: str
    document_type: str | None = None
    warnings: tuple[str, ...] = ()
    ocr_status: str | None = None
    document_date: datetime | None = None


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    document_type: str | None = None
    warnings: tuple[str, ...] = ()
    attachment_names: tuple[str, ...] = ()
    attachments: tuple[ExtractedAttachment, ...] = ()
    ocr_status: str | None = None
    sender: str | None = None
    recipients: str | None = None
    cc: str | None = None
    bcc: str | None = None
    subject: str | None = None
    document_date: datetime | None = None


def extract_document(path: Path, file_type: str) -> ExtractedDocument:
    return _extract_document(path, file_type, depth=0)


def _extract_document(path: Path, file_type: str, depth: int) -> ExtractedDocument:
    extension = file_type.lower().lstrip(".")
    if extension == "pdf":
        text, warnings, ocr_status = _extract_pdf(path)
        return ExtractedDocument(
            text=text,
            document_type=_classify_text(path.name, text, fallback="pdf"),
            warnings=tuple(warnings),
            ocr_status=ocr_status,
            document_date=_extract_first_date(text),
        )
    if extension == "docx":
        text, warnings = _extract_docx(path)
        return ExtractedDocument(
            text=text,
            document_type=_classify_text(path.name, text, fallback="document"),
            warnings=tuple(warnings),
            document_date=_extract_first_date(text),
        )
    if extension == "eml":
        return _extract_email(path, depth=depth)
    if extension == "html" or extension == "htm":
        text = _extract_html(_read_text(path))
        return ExtractedDocument(text=text, document_type=_classify_text(path.name, text, fallback="html"), document_date=_extract_first_date(text))
    if extension == "rtf":
        text = _extract_rtf(_read_text(path))
        return ExtractedDocument(text=text, document_type=_classify_text(path.name, text, fallback="rtf"), document_date=_extract_first_date(text))
    if extension == "xlsx":
        text, warnings = _extract_xlsx(path)
        return ExtractedDocument(text=text, document_type=_classify_text(path.name, text, fallback="spreadsheet"), warnings=tuple(warnings), document_date=_extract_first_date(text))
    if extension == "zip":
        text, attachment_names, warnings = _extract_zip_archive(path, depth)
        return ExtractedDocument(text=text, document_type="archive", attachment_names=tuple(attachment_names), warnings=tuple(warnings), document_date=_extract_first_date(text))
    if extension in {"txt", "md", "csv", "tsv", "log"}:
        text = _read_text(path)
        return ExtractedDocument(
            text=text,
            document_type=_classify_text(path.name, text),
            document_date=_extract_first_date(text),
        )
    return ExtractedDocument(text="", document_type=extension or None)


def _extract_pdf(path: Path) -> tuple[str, list[str], str | None]:
    warnings = []
    try:
        reader = PdfReader(str(path))
    except PdfReadError as error:
        return "", [f"PDF extraction failed: {error}"], "failed"
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            return "", ["PDF is encrypted and could not be decrypted."], "blocked_encrypted"

    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as error:
            warnings.append(f"Page {page_number} text extraction failed: {error}")
            page_text = ""
        if page_text.strip():
            pages.append(f"[Page {page_number}]\n{page_text.strip()}")
    text = "\n\n".join(pages)
    ocr_status = "recommended" if not text.strip() and len(reader.pages) > 0 else "not_required"
    if ocr_status == "recommended":
        ocr_text, ocr_warnings = _run_pdf_ocr(path)
        warnings.extend(ocr_warnings)
        if ocr_text.strip():
            return ocr_text.strip(), warnings, "completed"
        warnings.append("No extractable PDF text found; OCR is recommended.")
    return text, warnings, ocr_status


def _extract_docx(path: Path) -> tuple[str, list[str]]:
    try:
        document = DocxDocument(str(path))
    except Exception as error:
        return "", [f"DOCX extraction failed: {error}"]
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    sections = paragraphs + table_rows
    return "\n\n".join(sections), []


def _extract_email(path: Path, depth: int = 0) -> ExtractedDocument:
    with path.open("rb") as file:
        message = message_from_binary_file(file)
    body, attachment_names, attachments, warnings = _email_parts(message, depth=depth)
    attachment_texts = [f"[Attachment: {attachment.filename}]\n{attachment.text}" for attachment in attachments if attachment.text]
    text_sections = [body, *attachment_texts]
    return ExtractedDocument(
        text="\n\n".join(section.strip() for section in text_sections if section.strip()),
        document_type="email",
        warnings=tuple(warnings),
        attachment_names=tuple(attachment_names),
        attachments=tuple(attachments),
        sender=message.get("From"),
        recipients=message.get("To"),
        cc=message.get("Cc"),
        bcc=message.get("Bcc"),
        subject=message.get("Subject"),
        document_date=_normalize_datetime(_parse_email_date(message.get("Date"))),
    )


def _email_parts(message: Message, depth: int = 0) -> tuple[str, list[str], list[ExtractedAttachment], list[str]]:
    body_parts = []
    attachment_names = []
    attachments = []
    warnings = []
    if message.is_multipart():
        for part in message.walk():
            disposition = str(part.get("Content-Disposition", "")).lower()
            filename = part.get_filename()
            if filename:
                attachment_names.append(filename)
            payload = part.get_payload(decode=True)
            if not payload:
                continue
            content_type = part.get_content_type()
            if content_type == "text/plain" and "attachment" not in disposition:
                decoded = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
                body_parts.append(decoded)
            elif filename:
                extracted_attachment, extraction_warnings = _extract_attachment_payload(
                    filename=filename,
                    payload=payload,
                    content_charset=part.get_content_charset(),
                    depth=depth,
                )
                warnings.extend(extraction_warnings)
                if extracted_attachment:
                    attachments.append(extracted_attachment)
            elif "attachment" in disposition:
                warnings.append(f"Attachment not text-extracted: {filename or content_type}")
        return (
            "\n\n".join(part.strip() for part in body_parts if part.strip()),
            attachment_names,
            attachments,
            warnings,
        )

    payload = message.get_payload(decode=True)
    if payload is None:
        return str(message.get_payload() or ""), attachment_names, attachments, warnings
    return payload.decode(message.get_content_charset() or "utf-8", errors="replace"), attachment_names, attachments, warnings


def _read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_attachment_payload(
    *,
    filename: str,
    payload: bytes,
    content_charset: str | None,
    depth: int,
) -> tuple[ExtractedAttachment | None, list[str]]:
    if depth >= 2:
        return None, [f"Attachment recursion limit reached: {filename}"]
    extension = Path(filename).suffix.lower().lstrip(".")
    if not extension:
        return None, [f"Attachment not text-extracted: {filename}"]

    if extension in {"txt", "csv", "tsv", "md", "log"}:
        text = payload.decode(content_charset or "utf-8", errors="replace")
        return (
            ExtractedAttachment(
                filename=filename,
                file_type=extension,
                payload=payload,
                text=text,
                document_type=_classify_text(filename, text),
                document_date=_extract_first_date(text),
            ),
            [],
        )
    if extension not in {"pdf", "docx", "eml", "html", "htm", "rtf", "xlsx", "zip"}:
        return None, [f"Attachment not text-extracted: {filename}"]

    with tempfile.TemporaryDirectory(prefix="ediscovery_attachment_") as temp_dir:
        attachment_path = Path(temp_dir) / Path(filename).name
        attachment_path.write_bytes(payload)
        extracted = _extract_document(attachment_path, extension, depth=depth + 1)
    warnings = [f"{filename}: {warning}" for warning in extracted.warnings]
    return (
        ExtractedAttachment(
            filename=filename,
            file_type=extension,
            payload=payload,
            text=extracted.text,
            document_type=extracted.document_type,
            warnings=extracted.warnings,
            ocr_status=extracted.ocr_status,
            document_date=extracted.document_date,
        ),
        warnings,
    )


def _extract_html(text: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _extract_rtf(text: str) -> str:
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", text).strip()


def _extract_xlsx(path: Path) -> tuple[str, list[str]]:
    warnings = []
    try:
        with zipfile.ZipFile(path) as workbook:
            shared_strings = _xlsx_shared_strings(workbook)
            rows = []
            for name in sorted(workbook.namelist()):
                if not name.startswith("xl/worksheets/") or not name.endswith(".xml"):
                    continue
                rows.extend(_xlsx_sheet_rows(workbook.read(name), shared_strings))
            return "\n".join(row for row in rows if row.strip()), warnings
    except Exception as error:
        return "", [f"XLSX extraction failed: {error}"]


def _xlsx_shared_strings(workbook: zipfile.ZipFile) -> list[str]:
    try:
        root = ElementTree.fromstring(workbook.read("xl/sharedStrings.xml"))
    except KeyError:
        return []
    strings = []
    for item in root.iter():
        if item.tag.endswith("}si") or item.tag == "si":
            strings.append(" ".join(text.text or "" for text in item.iter() if text.tag.endswith("}t") or text.tag == "t"))
    return strings


def _xlsx_sheet_rows(sheet_xml: bytes, shared_strings: list[str]) -> list[str]:
    root = ElementTree.fromstring(sheet_xml)
    rows = []
    for row in root.iter():
        if not (row.tag.endswith("}row") or row.tag == "row"):
            continue
        values = []
        for cell in row:
            if not (cell.tag.endswith("}c") or cell.tag == "c"):
                continue
            cell_type = cell.attrib.get("t")
            value = next((child.text for child in cell if child.tag.endswith("}v") or child.tag == "v"), None)
            if value is None:
                continue
            if cell_type == "s":
                try:
                    values.append(shared_strings[int(value)])
                except (IndexError, ValueError):
                    values.append(value)
            else:
                values.append(value)
        if values:
            rows.append(" | ".join(values))
    return rows


def _extract_zip_archive(path: Path, depth: int) -> tuple[str, list[str], list[str]]:
    if depth >= 2:
        return "", [], [f"Archive recursion limit reached: {path.name}"]
    texts = []
    names = []
    warnings = []
    try:
        with zipfile.ZipFile(path) as archive:
            for member in archive.infolist():
                if member.is_dir():
                    continue
                filename = Path(member.filename).name
                if not filename:
                    continue
                names.append(filename)
                payload = archive.read(member)
                attachment, attachment_warnings = _extract_attachment_payload(
                    filename=filename,
                    payload=payload,
                    content_charset=None,
                    depth=depth,
                )
                warnings.extend([f"{filename}: {warning}" for warning in attachment_warnings])
                if attachment and attachment.text:
                    texts.append(f"[Archive item: {filename}]\n{attachment.text}")
    except Exception as error:
        return "", names, [f"ZIP extraction failed: {error}"]
    return "\n\n".join(texts), names, warnings


def _parse_email_date(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        return parsedate_to_datetime(value)
    except (TypeError, ValueError):
        return None


def _classify_text(filename: str, text: str, fallback: str = "text") -> str:
    lowered = f"{filename}\n{text[:1000]}".lower()
    if "response to request for production" in lowered or "request for production no." in lowered:
        return "discovery_response"
    if "legal memo" in lowered or "legal memorandum" in lowered:
        return "legal_memo"
    if "non-disclosure" in lowered or "nda" in lowered:
        return "nda"
    if "invoice" in lowered or "amount due" in lowered or re.search(r"\binv-\d+", lowered):
        return "invoice"
    if "court order" in lowered or lowered.startswith("order "):
        return "court_order"
    if "motion" in lowered or "memorandum of law" in lowered or "complaint" in lowered:
        return "pleading"
    if "slack export" in lowered or "chat export" in lowered:
        return "chat_export"
    if "agreement" in lowered or "contract" in lowered:
        return "contract"
    return fallback


def _extract_first_date(text: str) -> datetime | None:
    match = re.search(
        r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
        r"Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.?\s+\d{1,2},?\s+\d{4}\b",
        text,
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    for fmt in ("%B %d, %Y", "%b %d, %Y", "%B %d %Y", "%b %d %Y"):
        try:
            return _normalize_datetime(datetime.strptime(match.group(0).replace(".", ""), fmt))
        except ValueError:
            continue
    return None


def _normalize_datetime(value: datetime | None) -> datetime | None:
    if value is None:
        return None
    if value.tzinfo is None:
        return value.replace(tzinfo=UTC)
    return value.astimezone(UTC)


def _is_text_attachment(filename: str | None) -> bool:
    return bool(filename and Path(filename).suffix.lower() in {".txt", ".csv", ".tsv", ".md", ".log"})


def _run_pdf_ocr(path: Path) -> tuple[str, list[str]]:
    if not settings.ocr_enabled:
        return "", ["OCR is disabled; set OCR_ENABLED=true and OCR_PDF_TO_TEXT_COMMAND to enable scanned PDF OCR."]
    if not settings.ocr_pdf_to_text_command:
        return "", ["OCR is enabled but OCR_PDF_TO_TEXT_COMMAND is not configured."]

    command = [part.format(input=str(path)) for part in shlex.split(settings.ocr_pdf_to_text_command)]
    try:
        completed = subprocess.run(
            command,
            capture_output=True,
            check=False,
            text=True,
            timeout=settings.ocr_timeout_seconds,
        )
    except FileNotFoundError:
        return "", [f"OCR command not found: {command[0]}"]
    except subprocess.TimeoutExpired:
        return "", [f"OCR timed out after {settings.ocr_timeout_seconds} seconds."]

    warnings = []
    if completed.stderr.strip():
        warnings.append(f"OCR stderr: {completed.stderr.strip()}")
    if completed.returncode != 0:
        warnings.append(f"OCR command failed with exit code {completed.returncode}.")
        return "", warnings
    return completed.stdout, warnings
